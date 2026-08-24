# -*- coding: utf-8 -*-
"""docx 批注导出（可选产物 P7）：在作业文档副本中写入 Word 批注（w:comment）。

参考 GradingServer `annotate.py` 的 OOXML 注入方案，参数化并去除 grading 依赖：
- word/document.xml：目标段落前插 w:commentRangeStart、段末追加 w:commentReference run、段后 w:commentRangeEnd
- word/comments.xml：新建（w:comments → w:comment，author/date/批注文本）
- [Content_Types].xml：注册 comments 部件
- word/_rels/document.xml.rels：添加 comments 关系

只对未通过项（ok=False）生成批注。**不修改学生原始文件**，返回带批注的副本 bytes。
仅支持 docx；xlsx/pptx 批注机制不同（xlsx 有 Comment，pptx 无标准批注），本模块抛 ValueError。

用法：
    from annotate import annotate_docx
    annotated = annotate_docx(original_bytes, diffs)
    Path("学生_annotated.docx").write_bytes(annotated)
"""
import zipfile
from datetime import datetime
from io import BytesIO

from lxml import etree
# 用技能自己的 parser 拿标题段（解析最新、无 grading 依赖）
try:
    from parsers import parse_bytes  # noqa: F401
    _HAS_PARSERS = True
except ImportError:
    _HAS_PARSERS = False

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
_R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

COMMENT_AUTHOR = "评分系统"


def _w(tag):
    return "{%s}%s" % (_W, tag)


def _comment_text(d):
    """单条扣分批注文本：评分项 + 期望/实际 + 位置 + 分值 + 依据。"""
    label = d.get("label") or d.get("dim") or "评分项"
    loc = d.get("position") or "全文"
    actual = d.get("actual")
    if isinstance(actual, (set, tuple, list)):
        actual = ", ".join(str(x) for x in actual) or "(无)"
    ev = d.get("evidence") or ""
    return "维度 %s 未通过：期望 %s，实际 %s（%s，扣 %s 分）。依据：%s" % (
        label, d.get("expected"), actual, loc, d.get("weight"), ev)


def _title_indices_from_parsed(parsed):
    """从技能 ParsedDoc 的 paragraphs 判断标题段（style∈标题样式 或 文档首段）。"""
    styles = {"Title", "标题", "Heading"}
    styles |= {f"Heading {n}" for n in range(1, 10)}
    styles |= {f"标题 {n}" for n in range(1, 10)}
    idx = {i for i, p in enumerate(parsed.paragraphs)
           if p.get("style") in styles}
    if not idx and parsed.paragraphs:
        for i, p in enumerate(parsed.paragraphs):
            if p.get("text", "").strip():
                idx.add(i)
                break
    return idx


def annotate_docx(data, diffs):
    """返回带批注的 docx bytes；diffs 含 {label/position/expected/actual/weight/ok/evidence}。"""
    fails = [d for d in diffs if not d.get("ok")]
    src = BytesIO(data)
    if not zipfile.is_zipfile(src):
        raise ValueError("不是有效的 docx 容器")
    with zipfile.ZipFile(src) as zin:
        names = zin.namelist()
        if "word/document.xml" not in names:
            raise ValueError("非 docx：缺 word/document.xml")
        doc_xml = zin.read("word/document.xml")
        ctypes_xml = zin.read("[Content_Types].xml")
        rels_xml = (zin.read("word/_rels/document.xml.rels")
                    if "word/_rels/document.xml.rels" in names else None)
        others = {n: zin.read(n) for n in names
                  if n not in ("word/document.xml", "[Content_Types].xml",
                               "word/_rels/document.xml.rels")}
        old_comments_xml = others.pop("word/comments.xml", None)

    if not fails:
        # 无扣分项：原样重打包返回（保留旧批注）
        out = BytesIO()
        with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
            zout.writestr("word/document.xml", doc_xml)
            zout.writestr("[Content_Types].xml", ctypes_xml)
            if rels_xml is not None:
                zout.writestr("word/_rels/document.xml.rels", rels_xml)
            if old_comments_xml is not None:
                zout.writestr("word/comments.xml", old_comments_xml)
            for n, b in others.items():
                zout.writestr(n, b)
        return out.getvalue()

    # 标题集：优先用解析器的贴心标题判断（body_list = 非标题段）
    from docx import Document as _DocxDocument
    doc_obj = _DocxDocument(BytesIO(data))
    if _HAS_PARSERS:
        try:
            parsed = parse_bytes(data, "word")
            title_set = _title_indices_from_parsed(parsed)
        except Exception:
            title_set = set()
    else:
        title_set = set()
    body_list = [i for i in range(len(doc_obj.paragraphs)) if i not in title_set]
    if not body_list:
        body_list = list(range(len(doc_obj.paragraphs)))
    if not body_list:
        raise ValueError("文档无正文段落，无法定位批注")

    doc = etree.fromstring(doc_xml)
    body = doc.find(_w("body"))
    paras = body.findall(_w("p")) if body is not None else []
    if not paras:
        raise ValueError("文档无段落，无法定位批注")

    def _target_para_index(d):
        loc = d.get("position") or ""
        match = None
        import re
        m = re.search(r"第\s*(\d+)\s*段", str(loc))
        if m:
            match = int(m.group(1))
        elif str(loc).startswith("段落#"):
            try:
                match = int(str(loc)[len("段落#"):])
            except ValueError:
                match = 0
        if match and 1 <= match <= len(body_list):
            return body_list[match - 1]
        return body_list[0] if body_list else 0

    # comments.xml 合并旧批注、新批注 id 续排
    comments_el = etree.Element(_w("comments"))
    next_cid = 0
    if old_comments_xml is not None:
        old_root = etree.fromstring(old_comments_xml)
        ids = []
        for old_c in old_root.findall(_w("comment")):
            comments_el.append(old_c)
            try:
                ids.append(int(old_c.get(_w("id"))))
            except (TypeError, ValueError):
                pass
        if ids:
            next_cid = max(ids)
    for i, d in enumerate(fails):
        cid = str(next_cid + 1 + i)
        idx = _target_para_index(d)
        if not (0 <= idx < len(paras)):
            idx = body_list[-1] if body_list else 0
        p = paras[idx]
        p.addprevious(etree.fromstring(
            '<w:commentRangeStart xmlns:w="%s" w:id="%s"/>' % (_W, cid)))
        r = etree.SubElement(p, _w("r"))
        etree.SubElement(r, _w("commentReference")).set(_w("id"), cid)
        p.addnext(etree.fromstring(
            '<w:commentRangeEnd xmlns:w="%s" w:id="%s"/>' % (_W, cid)))
        c = etree.SubElement(comments_el, _w("comment"))
        c.set(_w("id"), cid)
        c.set(_w("author"), COMMENT_AUTHOR)
        c.set(_w("date"), datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ"))
        cp = etree.SubElement(c, _w("p"))
        cr = etree.SubElement(cp, _w("r"))
        t = etree.SubElement(cr, _w("t"))
        t.set("{%s}space" % "http://www.w3.org/XML/1998/namespace", "preserve")
        t.text = _comment_text(d)

    new_doc = etree.tostring(doc, xml_declaration=True, encoding="UTF-8", standalone=True)
    new_comments = etree.tostring(comments_el, xml_declaration=True,
                                  encoding="UTF-8", standalone=True)

    ct = etree.fromstring(ctypes_xml)
    has_override = any(
        o.get("PartName") == "/word/comments.xml"
        for o in ct.findall("{%s}Override" % _CT_NS))
    if not has_override:
        override = etree.SubElement(ct, "{%s}Override" % _CT_NS)
        override.set("PartName", "/word/comments.xml")
        override.set("ContentType",
                     "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml")
    new_ct = etree.tostring(ct, xml_declaration=True, encoding="UTF-8", standalone=True)

    if rels_xml is None:
        rels = etree.fromstring('<Relationships xmlns="%s"/>' % _R_NS)
    else:
        rels = etree.fromstring(rels_xml)
    has_comments_rel = any(
        r.get("Type") and "comments" in r.get("Type")
        for r in rels.findall("{%s}Relationship" % _R_NS))
    if not has_comments_rel:
        rel = etree.SubElement(rels, "{%s}Relationship" % _R_NS)
        rel.set("Id", "rIdComments")
        rel.set("Type",
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments")
        rel.set("Target", "comments.xml")
    new_rels = etree.tostring(rels, xml_declaration=True, encoding="UTF-8", standalone=True)

    out = BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        zout.writestr("word/document.xml", new_doc)
        zout.writestr("word/comments.xml", new_comments)
        zout.writestr("[Content_Types].xml", new_ct)
        zout.writestr("word/_rels/document.xml.rels", new_rels)
        for n, b in others.items():
            zout.writestr(n, b)
    return out.getvalue()


def annotate(parsed_bytes, diffs, file_type="word"):
    """统一入口：word→docx 批注；xlsx/pptx 暂不支持（抛 ValueError）。"""
    if file_type != "word":
        raise ValueError("批注导出当前仅支持 word（docx）；xlsx/pptx 批注机制不同，请另行评估")
    return annotate_docx(parsed_bytes, diffs)


if __name__ == "__main__":
    import sys
    from pathlib import Path

    diffs_file, out_file = sys.argv[1], sys.argv[2]
    inp = Path(diffs_file).resolve()
    # 期望 diffs.json 结构：{"source": "<原始docx路径>", "diffs": [...]}
    import json
    meta = json.loads(inp.read_text(encoding="utf-8"))
    src = Path(meta["source"]).resolve()
    annotated = annotate_docx(src.read_bytes(), meta["diffs"])
    Path(out_file).write_bytes(annotated)
    print("已生成带批注副本:", out_file)
